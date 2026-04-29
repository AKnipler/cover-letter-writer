
import re
import regex as re
import json
import os
import json
from docx import Document
from python_docx_replace import docx_replace
from docx2pdf import convert
import time
from datetime import datetime
from bs4 import BeautifulSoup
import shutil
import requests
from html import unescape




def extract_job_info(html):
    # Ensure html is text for regex operations (some callers pass bytes)
    if isinstance(html, (bytes, bytearray)):
        html = html.decode("utf-8", errors="replace")

    soup = BeautifulSoup(html, "html.parser")

    # 1. JSON (best for SEEK, modern apps)
    company, title = extract_from_json(html)

    # 2. JSON-LD (LinkedIn, Indeed, etc.)
    if not company or not title:
        c2, t2 = extract_from_json_ld(soup)
        company = company or c2
        title = title or t2

    # 3. Meta/title fallback (GradConnection etc.)
    if not company or not title:
        c3, t3 = extract_from_meta(soup)
        company = company or c3
        title = title or t3

    return (company, title)

def extract_from_json(html):
    """
    Extract from raw JSON blobs inside script tags or inline JS
    """
    company = None
    title = None

    # Common patterns
    title_patterns = [
        r'"title":"([^"]+)"',
        r'"normalisedRoleTitle":"([^"]+)"'
    ]

    for pattern in title_patterns:
        title_match = re.search(pattern, html)
        if title_match:
            title = title_match.group(1)
            break

    # Try multiple company patterns (important!)
    company_patterns = [
        r'"advertiser":\{"name":"([^"]+)"',
        r'"companyProfile":\{"name":"([^"]+)"',
        r'"normalisedOrganisationName":"([^"]+)"'
    ]

    for pattern in company_patterns:
        match = re.search(pattern, html)
        if match:
            company = match.group(1)
            break

    return company, title

def extract_from_json_ld(soup):
    """
    Extract from structured data (application/ld+json)
    """
    scripts = soup.find_all("script", type="application/ld+json")

    for script in scripts:
        try:
            if not script.string:
                continue

            data = json.loads(script.string)

            # Handle @graph structure
            if isinstance(data, dict) and "@graph" in data:
                items = data["@graph"]
            elif isinstance(data, list):
                items = data
            else:
                items = [data]

            for item in items:
                title = item.get("title") or item.get("jobTitle")

                company = None
                if "hiringOrganization" in item:
                    company = item["hiringOrganization"].get("name")

                if title or company:
                    return company, title

        except Exception:
            continue

    return None, None


def extract_from_meta(soup):
    """
    Extract from og:title or <title>
    """
    og = soup.find("meta", property="og:title")

    if og and og.get("content"):
        main = og["content"].split("|")[0].strip()

        # SEEK format: "Job Title Job in Location - SEEK"
        if " Job in " in main:
            title = main.split(" Job in ")[0]
            return None, title

    # Meta description (extract company)
    desc = soup.find("meta", attrs={"name": "description"})
    if desc and desc.get("content"):
        text = desc["content"]

        # crude but effective
        match = re.search(r'with\s+([A-Z][A-Za-z0-9& ]+)', text)
        if match:
            company = match.group(1).strip()
            return company, None

    return None, None





def ordinal_suffix(day):
    """Adds an English ordinal suffix to the day of the month."""
    if 4 <= day % 100 <= 20:
        suffix = "th"
    else:
        suffix = {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
    return f"{suffix}"

def create_cover_letter(template_path, output_path, cover_letter_content):
    
    # To copy to a specific new file name
    try:
        shutil.copy2(template_path, output_path)
        print(f"File copied to: {output_path}")
    except shutil.SameFileError:
        print("Source and destination represent the same file.")
    except PermissionError:
        print("Permission denied.")
    except FileNotFoundError:
        print("Source file not found.")
    except Exception as e:
        print(f"An error occurred: {e}")
    
    # Replace the placeholder text with the generated cover letter content
    for replacement in cover_letter_content:
        replace_text_in_docx(output_path, replacement, cover_letter_content[replacement])

    print(f"Text replaced and saved to {output_path}")



def replace_text_in_docx(doc_path, old_text, new_text):
    # Load the document
    doc = Document(doc_path)
    replaced = False
    
    # Iterate through all paragraphs in the document
    for p in doc.paragraphs:
        for i in range(len(p.runs) - 1, -1, -1):

            if "]" in p.runs[i].text and "[" not in p.runs[i].text:
                p.runs[i-1].add_text(p.runs[i].text)
                p.runs[i].clear()
                continue

            if old_text in p.runs[i].text:
                # Replacing text in the whole paragraph will remove original formatting
                
                if "DATE" in old_text:
                    date1 = datetime.now().strftime("%A, %#d") # replace %#d with %-d for platforms that don't support it (e.g. Linux)
                    date2 = ordinal_suffix(datetime.now().day)
                    date3 = datetime.now().strftime(" of %B %Y")

                    p.runs[i].text = p.runs[i].text.replace(old_text, date1)
                    suffix = p.add_run(date2)
                    # Font formatting for the ordinal suffix
                    suffix.font.color.rgb = p.runs[i].font.color.rgb
                    suffix.font.size = p.runs[i].font.size
                    suffix.font.name = p.runs[i].font.name
                    suffix.style = p.runs[i].style
                    suffix.font.superscript = True

                    # Font formatting for the rest of the date
                    trail = p.add_run(date3)
                    trail.font.color.rgb = p.runs[i].font.color.rgb
                    trail.font.size = p.runs[i].font.size
                    trail.font.name = p.runs[i].font.name
                    trail.style = p.runs[i].style
                    
                    replaced = True

                else:
                    # Save the current formatting of the run
                    temp = []
                    temp.append(p.runs[i].font.color.rgb)
                    temp.append(p.runs[i].font.size)
                    temp.append(p.runs[i].font.name)
                    temp.append(p.runs[i].font.bold)
                    temp.append(p.runs[i].style)

                    p.runs[i].text = p.runs[i].text.replace(old_text, new_text)

                    # Restore the original formatting
                    p.runs[i].font.color.rgb = temp[0]
                    p.runs[i].font.size = temp[1]
                    p.runs[i].font.name = temp[2]
                    p.runs[i].font.bold = temp[3]
                    p.runs[i].style = temp[4]

    # Save the modified document
    doc.save(doc_path)



def send_request(URL):
    response = requests.get(
            url='https://proxy.scrapeops.io/v1/',
            params={
                'api_key': 'c0243109-cacb-4fcd-add5-68c646ef3508',
                'url': URL,
                'render_js': 'true'
            }
          )

    
    # print('Response HTTP Status Code: ', response.status_code)
    print('Response HTTP Response Body: ', response.content)

    return response.content
  
def page_extraction(content):
    
    # Use BeautifulSoup to parse the HTML content and extract the text
    soup = BeautifulSoup(content, 'html.parser')

    # Find the script containing window.SK_DL
    # script_text = soup.find_all("script")

    # target_script = None
    # for script in script_text:
    #     if script.string and "window.SK_DL" in script.string:
    #         target_script = script.string
    #         break

    # # Extract JSON part
    # match = re.search(r'window\.SK_DL\s*=\s*({.*?});', target_script, re.DOTALL)
    # data = json.loads(match.group(1))

    # advertiser_name = data["advertiserName"]
    # position_title = data["jobTitle"]

    (advertiser_name, position_title) = extract_job_info(content)

    # for tag in soup(["script", "style"]):
    #     tag.decompose()

    # for tag in soup.select(".hidden"):
    #     tag.decompose()

    script = soup.find("script", string=re.compile("SEEK_REDUX_DATA"))
    script_text = script.string

    match = re.search(r'window\.SEEK_REDUX_DATA\s*=\s*(\{.*\});', script_text, re.DOTALL)
    
    
    start = script_text.find("window.SEEK_REDUX_DATA =") 
    start = script_text.find("{", start)

    brace_count = 0
    end = start

    for i, char in enumerate(script_text[start:]):
        if char == "{":
            brace_count += 1
        elif char == "}":
            brace_count -= 1
            if brace_count == 0:
                end = start + i + 1
                break

    json_str = script_text[start:end]

    data = json.loads(json_str)

    print('Extracted Data: ', data)

    # body = body.encode('latin1').decode('utf-8')
    content_html = data["jobdetails"]["result"]["job"]["content"]
    decoded = unescape(content_html)

    clean_body = BeautifulSoup(decoded, "html.parser").get_text(" ", strip=True)

    print('Encoded Content: ', clean_body)
    

    return clean_body, advertiser_name, position_title



